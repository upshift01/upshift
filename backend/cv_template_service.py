"""
CV Template Service - Handles .docx template upload, management, and document generation
Supports placeholder replacement and both .docx and PDF output formats
"""
import os
import re
import io
import logging
import shutil
import subprocess
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional
from uuid import uuid4
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Template storage paths
TEMPLATE_STORAGE_PATH = "/app/public/uploads/cv_templates"
GENERATED_CV_PATH = "/app/public/uploads/generated_cvs"

# Ensure directories exist
os.makedirs(TEMPLATE_STORAGE_PATH, exist_ok=True)
os.makedirs(GENERATED_CV_PATH, exist_ok=True)

# Template categories
TEMPLATE_CATEGORIES = [
    {"id": "professional", "name": "Professional", "description": "Clean and corporate styles"},
    {"id": "creative", "name": "Creative", "description": "Modern and artistic designs"},
    {"id": "ats-friendly", "name": "ATS-Friendly", "description": "Optimized for applicant tracking systems"},
    {"id": "academic", "name": "Academic", "description": "For education and research roles"},
    {"id": "executive", "name": "Executive", "description": "For senior leadership positions"},
    {"id": "entry-level", "name": "Entry Level", "description": "For graduates and early career"},
    {"id": "industry-specific", "name": "Industry Specific", "description": "Tailored for specific industries"},
    {"id": "minimalist", "name": "Minimalist", "description": "Simple and clean designs"},
]

# Standard placeholders that templates should use
PLACEHOLDER_MAP = {
    # Personal Information
    "{{FULL_NAME}}": "full_name",
    "{{EMAIL}}": "email",
    "{{PHONE}}": "phone",
    "{{ADDRESS}}": "address",
    "{{LOCATION}}": "address",
    "{{ID_NUMBER}}": "id_number",
    "{{PHOTO}}": "photo",
    "{{LANGUAGES}}": "languages",
    "{{LINKEDIN_URL}}": "linkedin_url",
    "{{LINKEDIN}}": "linkedin_url",
    
    # Professional Summary
    "{{SUMMARY}}": "summary",
    "{{PROFESSIONAL_SUMMARY}}": "summary",
    "{{PROFILE}}": "summary",
    "{{OBJECTIVE}}": "summary",
    
    # Skills
    "{{SKILLS}}": "skills",
    "{{SKILL_LIST}}": "skills",
    
    # Experience section markers
    "{{EXPERIENCE_SECTION}}": "experiences",
    "{{WORK_EXPERIENCE}}": "experiences",
    
    # Education section markers
    "{{EDUCATION_SECTION}}": "education",
    
    # References section marker
    "{{REFERENCES_SECTION}}": "references",
    "{{REFERENCES}}": "references",
    
    # Individual experience placeholders (for templates with fixed slots)
    "{{EXP_1_TITLE}}": "experiences[0].title",
    "{{EXP_1_COMPANY}}": "experiences[0].company",
    "{{EXP_1_DURATION}}": "experiences[0].duration",
    "{{EXP_1_DESCRIPTION}}": "experiences[0].description",
    "{{EXP_1_ACHIEVEMENTS}}": "experiences[0].achievements",
    
    "{{EXP_2_TITLE}}": "experiences[1].title",
    "{{EXP_2_COMPANY}}": "experiences[1].company",
    "{{EXP_2_DURATION}}": "experiences[1].duration",
    "{{EXP_2_DESCRIPTION}}": "experiences[1].description",
    "{{EXP_2_ACHIEVEMENTS}}": "experiences[1].achievements",
    
    "{{EXP_3_TITLE}}": "experiences[2].title",
    "{{EXP_3_COMPANY}}": "experiences[2].company",
    "{{EXP_3_DURATION}}": "experiences[2].duration",
    "{{EXP_3_DESCRIPTION}}": "experiences[2].description",
    "{{EXP_3_ACHIEVEMENTS}}": "experiences[2].achievements",
    
    # Individual education placeholders
    "{{EDU_1_DEGREE}}": "education[0].degree",
    "{{EDU_1_INSTITUTION}}": "education[0].institution",
    "{{EDU_1_YEAR}}": "education[0].year",
    
    "{{EDU_2_DEGREE}}": "education[1].degree",
    "{{EDU_2_INSTITUTION}}": "education[1].institution",
    "{{EDU_2_YEAR}}": "education[1].year",
    
    # Individual reference placeholders (for templates with fixed slots)
    "{{REF_1_NAME}}": "references[0].name",
    "{{REF_1_TITLE}}": "references[0].title",
    "{{REF_1_COMPANY}}": "references[0].company",
    "{{REF_1_EMAIL}}": "references[0].email",
    "{{REF_1_PHONE}}": "references[0].phone",
    
    "{{REF_2_NAME}}": "references[1].name",
    "{{REF_2_TITLE}}": "references[1].title",
    "{{REF_2_COMPANY}}": "references[1].company",
    "{{REF_2_EMAIL}}": "references[1].email",
    "{{REF_2_PHONE}}": "references[1].phone",
    
    "{{REF_3_NAME}}": "references[2].name",
    "{{REF_3_TITLE}}": "references[2].title",
    "{{REF_3_COMPANY}}": "references[2].company",
    "{{REF_3_EMAIL}}": "references[2].email",
    "{{REF_3_PHONE}}": "references[2].phone",
    
    # Certifications section marker
    "{{CERTIFICATIONS_SECTION}}": "certifications",
    "{{CERTIFICATIONS}}": "certifications",
    
    # Individual certification placeholders
    "{{CERT_1_NAME}}": "certifications[0].name",
    "{{CERT_1_ORGANIZATION}}": "certifications[0].organization",
    "{{CERT_1_ISSUE_DATE}}": "certifications[0].issueDate",
    "{{CERT_1_EXPIRY_DATE}}": "certifications[0].expiryDate",
    "{{CERT_1_CREDENTIAL_ID}}": "certifications[0].credentialId",
    "{{CERT_1_URL}}": "certifications[0].url",
    
    "{{CERT_2_NAME}}": "certifications[1].name",
    "{{CERT_2_ORGANIZATION}}": "certifications[1].organization",
    "{{CERT_2_ISSUE_DATE}}": "certifications[1].issueDate",
    "{{CERT_2_EXPIRY_DATE}}": "certifications[1].expiryDate",
    "{{CERT_2_CREDENTIAL_ID}}": "certifications[1].credentialId",
    "{{CERT_2_URL}}": "certifications[1].url",
    
    "{{CERT_3_NAME}}": "certifications[2].name",
    "{{CERT_3_ORGANIZATION}}": "certifications[2].organization",
    "{{CERT_3_ISSUE_DATE}}": "certifications[2].issueDate",
    "{{CERT_3_EXPIRY_DATE}}": "certifications[2].expiryDate",
    "{{CERT_3_CREDENTIAL_ID}}": "certifications[2].credentialId",
    "{{CERT_3_URL}}": "certifications[2].url",
}


class CVTemplateService:
    def __init__(self, db):
        self.db = db
    
    async def upload_template(
        self,
        file_content: bytes,
        filename: str,
        template_name: str,
        category: str,
        description: str = "",
        uploaded_by: str = "admin",
        reseller_id: Optional[str] = None,
        preview_image: Optional[bytes] = None
    ) -> Dict[str, Any]:
        """
        Upload a new .docx CV template
        """
        try:
            # Validate file is a docx
            if not filename.lower().endswith('.docx'):
                raise ValueError("Only .docx files are allowed")
            
            # Generate unique ID and filename
            template_id = str(uuid4())
            safe_filename = f"{template_id}.docx"
            
            # Determine storage path (platform or reseller-specific)
            if reseller_id:
                storage_dir = os.path.join(TEMPLATE_STORAGE_PATH, "resellers", reseller_id)
            else:
                storage_dir = os.path.join(TEMPLATE_STORAGE_PATH, "platform")
            
            os.makedirs(storage_dir, exist_ok=True)
            file_path = os.path.join(storage_dir, safe_filename)
            
            # Save the template file
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            # Validate template has placeholders
            placeholders_found = self._extract_placeholders(file_content)
            
            # Save preview image if provided
            preview_url = None
            if preview_image:
                preview_filename = f"{template_id}_preview.png"
                preview_path = os.path.join(storage_dir, preview_filename)
                with open(preview_path, 'wb') as f:
                    f.write(preview_image)
                preview_url = f"/uploads/cv_templates/{'resellers/' + reseller_id if reseller_id else 'platform'}/{preview_filename}"
            
            # Create template record
            template_record = {
                "id": template_id,
                "name": template_name,
                "filename": safe_filename,
                "original_filename": filename,
                "category": category,
                "description": description,
                "file_path": file_path,
                "file_url": f"/uploads/cv_templates/{'resellers/' + reseller_id if reseller_id else 'platform'}/{safe_filename}",
                "preview_url": preview_url,
                "placeholders_found": placeholders_found,
                "uploaded_by": uploaded_by,
                "reseller_id": reseller_id,
                "is_platform_template": reseller_id is None,
                "is_active": True,
                "usage_count": 0,
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }
            
            await self.db.cv_templates.insert_one(template_record)
            
            logger.info(f"CV template uploaded: {template_name} (ID: {template_id})")
            
            return {
                "success": True,
                "template_id": template_id,
                "name": template_name,
                "category": category,
                "placeholders_found": placeholders_found,
                "message": f"Template '{template_name}' uploaded successfully"
            }
            
        except Exception as e:
            logger.error(f"Error uploading CV template: {str(e)}")
            raise
    
    def _extract_placeholders(self, file_content: bytes) -> List[str]:
        """Extract all placeholders from a docx template"""
        try:
            doc = Document(io.BytesIO(file_content))
            placeholders = set()
            
            # Check paragraphs
            for para in doc.paragraphs:
                found = re.findall(r'\{\{[A-Z0-9_]+\}\}', para.text)
                placeholders.update(found)
            
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            found = re.findall(r'\{\{[A-Z0-9_]+\}\}', para.text)
                            placeholders.update(found)
            
            return list(placeholders)
        except Exception as e:
            logger.warning(f"Error extracting placeholders: {str(e)}")
            return []
    
    async def get_templates(
        self,
        category: Optional[str] = None,
        reseller_id: Optional[str] = None,
        include_platform: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Get available CV templates
        """
        query = {"is_active": True}
        
        if reseller_id:
            if include_platform:
                query["$or"] = [
                    {"reseller_id": reseller_id},
                    {"is_platform_template": True}
                ]
            else:
                query["reseller_id"] = reseller_id
        else:
            query["is_platform_template"] = True
        
        if category:
            query["category"] = category
        
        templates = await self.db.cv_templates.find(
            query,
            {"_id": 0}
        ).sort("name", 1).to_list(100)
        
        return templates
    
    async def get_template_by_id(self, template_id: str) -> Optional[Dict[str, Any]]:
        """Get a specific template by ID"""
        return await self.db.cv_templates.find_one(
            {"id": template_id, "is_active": True},
            {"_id": 0}
        )
    
    async def delete_template(self, template_id: str, deleted_by: str) -> bool:
        """Soft delete a template"""
        result = await self.db.cv_templates.update_one(
            {"id": template_id},
            {
                "$set": {
                    "is_active": False,
                    "deleted_at": datetime.now(timezone.utc),
                    "deleted_by": deleted_by
                }
            }
        )
        return result.modified_count > 0
    
    async def generate_cv_from_template(
        self,
        template_id: str,
        cv_data: Dict[str, Any],
        output_format: str = "pdf",  # "pdf" or "docx"
        user_id: str = None
    ) -> Dict[str, Any]:
        """
        Generate a CV from a template with user data
        """
        try:
            # Get template
            template = await self.get_template_by_id(template_id)
            if not template:
                raise ValueError(f"Template not found: {template_id}")
            
            # Load template document
            doc = Document(template["file_path"])
            
            # Prepare data for replacement
            replacement_data = self._prepare_replacement_data(cv_data)
            
            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                self._replace_in_paragraph(para, replacement_data)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, replacement_data)
            
            # Generate unique filename
            output_id = str(uuid4())
            user_name = cv_data.get("full_name", "CV").replace(" ", "_")
            
            # Save the filled document
            docx_filename = f"{output_id}_{user_name}.docx"
            docx_path = os.path.join(GENERATED_CV_PATH, docx_filename)
            doc.save(docx_path)
            
            result = {
                "success": True,
                "document_id": output_id,
                "template_used": template["name"],
                "format": output_format
            }
            
            if output_format == "pdf":
                # Convert to PDF
                pdf_filename = f"{output_id}_{user_name}.pdf"
                pdf_path = os.path.join(GENERATED_CV_PATH, pdf_filename)
                
                # Use LibreOffice for conversion (more reliable in server environment)
                conversion_success = self._convert_to_pdf(docx_path, pdf_path)
                
                if conversion_success and os.path.exists(pdf_path):
                    result["file_path"] = pdf_path
                    result["file_url"] = f"/uploads/generated_cvs/{pdf_filename}"
                    result["filename"] = pdf_filename
                    
                    # Read PDF content for base64 response
                    with open(pdf_path, 'rb') as f:
                        import base64
                        result["pdf_base64"] = base64.b64encode(f.read()).decode('utf-8')
                else:
                    # Fallback to docx if PDF conversion fails
                    logger.warning("PDF conversion failed, returning docx")
                    result["format"] = "docx"
                    result["file_path"] = docx_path
                    result["file_url"] = f"/uploads/generated_cvs/{docx_filename}"
                    result["filename"] = docx_filename
            else:
                result["file_path"] = docx_path
                result["file_url"] = f"/uploads/generated_cvs/{docx_filename}"
                result["filename"] = docx_filename
                
                # Read docx content for base64 response
                with open(docx_path, 'rb') as f:
                    import base64
                    result["docx_base64"] = base64.b64encode(f.read()).decode('utf-8')
            
            # Update template usage count
            await self.db.cv_templates.update_one(
                {"id": template_id},
                {"$inc": {"usage_count": 1}}
            )
            
            logger.info(f"CV generated from template {template_id} for user {user_id}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error generating CV from template: {str(e)}")
            raise
    
    def _prepare_replacement_data(self, cv_data: Dict[str, Any]) -> Dict[str, str]:
        """Prepare data for placeholder replacement"""
        replacements = {}
        
        # Simple fields
        replacements["{{FULL_NAME}}"] = cv_data.get("full_name", "")
        replacements["{{EMAIL}}"] = cv_data.get("email", "")
        replacements["{{PHONE}}"] = cv_data.get("phone", "")
        replacements["{{ADDRESS}}"] = cv_data.get("address", "")
        replacements["{{LOCATION}}"] = cv_data.get("address", "")
        replacements["{{ID_NUMBER}}"] = cv_data.get("id_number", "")
        replacements["{{LINKEDIN_URL}}"] = cv_data.get("linkedin_url", "")
        replacements["{{LINKEDIN}}"] = cv_data.get("linkedin_url", "")
        replacements["{{SUMMARY}}"] = cv_data.get("summary", "")
        replacements["{{PROFESSIONAL_SUMMARY}}"] = cv_data.get("summary", "")
        replacements["{{PROFILE}}"] = cv_data.get("summary", "")
        replacements["{{OBJECTIVE}}"] = cv_data.get("summary", "")
        
        # Languages
        languages = cv_data.get("languages", [])
        if isinstance(languages, list):
            # Handle both string and dict formats
            lang_strings = []
            for l in languages:
                if isinstance(l, dict):
                    lang_name = l.get("language", "")
                    proficiency = l.get("proficiency", "")
                    if lang_name:
                        if proficiency:
                            lang_strings.append(f"{lang_name} ({proficiency})")
                        else:
                            lang_strings.append(lang_name)
                elif isinstance(l, str) and l:
                    lang_strings.append(l)
            replacements["{{LANGUAGES}}"] = ", ".join(lang_strings)
        else:
            replacements["{{LANGUAGES}}"] = str(languages)
        
        # Skills
        skills = cv_data.get("skills", [])
        if isinstance(skills, list):
            replacements["{{SKILLS}}"] = ", ".join([s for s in skills if s])
            replacements["{{SKILL_LIST}}"] = "\n• ".join([s for s in skills if s])
            if replacements["{{SKILL_LIST}}"]:
                replacements["{{SKILL_LIST}}"] = "• " + replacements["{{SKILL_LIST}}"]
        else:
            replacements["{{SKILLS}}"] = str(skills)
            replacements["{{SKILL_LIST}}"] = str(skills)
        
        # Experience entries
        experiences = cv_data.get("experiences", [])
        for i, exp in enumerate(experiences[:5]):  # Support up to 5 experiences
            idx = i + 1
            replacements[f"{{{{EXP_{idx}_TITLE}}}}"] = exp.get("title", "")
            replacements[f"{{{{EXP_{idx}_COMPANY}}}}"] = exp.get("company", "")
            replacements[f"{{{{EXP_{idx}_DURATION}}}}"] = exp.get("duration", "")
            replacements[f"{{{{EXP_{idx}_DESCRIPTION}}}}"] = exp.get("description", "")
            achievements = exp.get("achievements", "")
            if isinstance(achievements, list):
                achievements = "\n• ".join(achievements)
            replacements[f"{{{{EXP_{idx}_ACHIEVEMENTS}}}}"] = achievements
        
        # Fill remaining experience slots with empty strings
        for i in range(len(experiences) + 1, 6):
            replacements[f"{{{{EXP_{i}_TITLE}}}}"] = ""
            replacements[f"{{{{EXP_{i}_COMPANY}}}}"] = ""
            replacements[f"{{{{EXP_{i}_DURATION}}}}"] = ""
            replacements[f"{{{{EXP_{i}_DESCRIPTION}}}}"] = ""
            replacements[f"{{{{EXP_{i}_ACHIEVEMENTS}}}}"] = ""
        
        # Experience section (formatted block)
        exp_section = []
        for exp in experiences:
            if exp.get("title") or exp.get("company"):
                exp_text = f"{exp.get('title', '')} at {exp.get('company', '')}\n"
                exp_text += f"{exp.get('duration', '')}\n"
                exp_text += f"{exp.get('description', '')}"
                if exp.get("achievements"):
                    exp_text += f"\nKey Achievements: {exp.get('achievements', '')}"
                exp_section.append(exp_text)
        replacements["{{EXPERIENCE_SECTION}}"] = "\n\n".join(exp_section)
        replacements["{{WORK_EXPERIENCE}}"] = replacements["{{EXPERIENCE_SECTION}}"]
        
        # Education entries
        education = cv_data.get("education", [])
        for i, edu in enumerate(education[:3]):  # Support up to 3 education entries
            idx = i + 1
            replacements[f"{{{{EDU_{idx}_DEGREE}}}}"] = edu.get("degree", "")
            replacements[f"{{{{EDU_{idx}_INSTITUTION}}}}"] = edu.get("institution", "")
            replacements[f"{{{{EDU_{idx}_YEAR}}}}"] = edu.get("year", "")
        
        # Fill remaining education slots
        for i in range(len(education) + 1, 4):
            replacements[f"{{{{EDU_{i}_DEGREE}}}}"] = ""
            replacements[f"{{{{EDU_{i}_INSTITUTION}}}}"] = ""
            replacements[f"{{{{EDU_{i}_YEAR}}}}"] = ""
        
        # Education section (formatted block)
        edu_section = []
        for edu in education:
            if edu.get("degree") or edu.get("institution"):
                edu_text = f"{edu.get('degree', '')}\n"
                edu_text += f"{edu.get('institution', '')} ({edu.get('year', '')})"
                edu_section.append(edu_text)
        replacements["{{EDUCATION_SECTION}}"] = "\n\n".join(edu_section)
        
        # References entries
        references = cv_data.get("references", [])
        for i, ref in enumerate(references[:3]):  # Support up to 3 references
            idx = i + 1
            replacements[f"{{{{REF_{idx}_NAME}}}}"] = ref.get("name", "")
            replacements[f"{{{{REF_{idx}_TITLE}}}}"] = ref.get("title", "")
            replacements[f"{{{{REF_{idx}_COMPANY}}}}"] = ref.get("company", "")
            replacements[f"{{{{REF_{idx}_EMAIL}}}}"] = ref.get("email", "")
            replacements[f"{{{{REF_{idx}_PHONE}}}}"] = ref.get("phone", "")
        
        # Fill remaining reference slots
        for i in range(len(references) + 1, 4):
            replacements[f"{{{{REF_{i}_NAME}}}}"] = ""
            replacements[f"{{{{REF_{i}_TITLE}}}}"] = ""
            replacements[f"{{{{REF_{i}_COMPANY}}}}"] = ""
            replacements[f"{{{{REF_{i}_EMAIL}}}}"] = ""
            replacements[f"{{{{REF_{i}_PHONE}}}}"] = ""
        
        # References section (formatted block)
        ref_section = []
        for ref in references:
            if ref.get("name"):
                ref_text = f"{ref.get('name', '')}"
                if ref.get("title"):
                    ref_text += f", {ref.get('title', '')}"
                if ref.get("company"):
                    ref_text += f" at {ref.get('company', '')}"
                contact_parts = []
                if ref.get("email"):
                    contact_parts.append(f"Email: {ref.get('email', '')}")
                if ref.get("phone"):
                    contact_parts.append(f"Phone: {ref.get('phone', '')}")
                if contact_parts:
                    ref_text += f"\n{' | '.join(contact_parts)}"
                ref_section.append(ref_text)
        replacements["{{REFERENCES_SECTION}}"] = "\n\n".join(ref_section)
        replacements["{{REFERENCES}}"] = replacements["{{REFERENCES_SECTION}}"]
        
        # Certifications entries
        certifications = cv_data.get("certifications", [])
        for i, cert in enumerate(certifications[:3]):  # Support up to 3 certifications
            idx = i + 1
            replacements[f"{{{{CERT_{idx}_NAME}}}}"] = cert.get("name", "")
            replacements[f"{{{{CERT_{idx}_ORGANIZATION}}}}"] = cert.get("organization", "")
            replacements[f"{{{{CERT_{idx}_ISSUE_DATE}}}}"] = cert.get("issueDate", "")
            replacements[f"{{{{CERT_{idx}_EXPIRY_DATE}}}}"] = cert.get("expiryDate", "")
            replacements[f"{{{{CERT_{idx}_CREDENTIAL_ID}}}}"] = cert.get("credentialId", "")
            replacements[f"{{{{CERT_{idx}_URL}}}}"] = cert.get("url", "")
        
        # Fill remaining certification slots
        for i in range(len(certifications) + 1, 4):
            replacements[f"{{{{CERT_{i}_NAME}}}}"] = ""
            replacements[f"{{{{CERT_{i}_ORGANIZATION}}}}"] = ""
            replacements[f"{{{{CERT_{i}_ISSUE_DATE}}}}"] = ""
            replacements[f"{{{{CERT_{i}_EXPIRY_DATE}}}}"] = ""
            replacements[f"{{{{CERT_{i}_CREDENTIAL_ID}}}}"] = ""
            replacements[f"{{{{CERT_{i}_URL}}}}"] = ""
        
        # Certifications section (formatted block)
        cert_section = []
        for cert in certifications:
            if cert.get("name"):
                cert_text = f"{cert.get('name', '')}"
                if cert.get("organization"):
                    cert_text += f" - {cert.get('organization', '')}"
                date_parts = []
                if cert.get("issueDate"):
                    date_parts.append(f"Issued: {cert.get('issueDate', '')}")
                if cert.get("expiryDate"):
                    date_parts.append(f"Expires: {cert.get('expiryDate', '')}")
                if date_parts:
                    cert_text += f"\n{' | '.join(date_parts)}"
                if cert.get("credentialId"):
                    cert_text += f"\nCredential ID: {cert.get('credentialId', '')}"
                if cert.get("url"):
                    cert_text += f"\nVerify: {cert.get('url', '')}"
                cert_section.append(cert_text)
        replacements["{{CERTIFICATIONS_SECTION}}"] = "\n\n".join(cert_section)
        replacements["{{CERTIFICATIONS}}"] = replacements["{{CERTIFICATIONS_SECTION}}"]
        
        return replacements
    
    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """Replace placeholders in a paragraph while preserving formatting"""
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # Handle run-level replacement to preserve formatting
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                
                # Also check full paragraph text (in case placeholder spans runs)
                if placeholder in paragraph.text:
                    inline = paragraph.runs
                    for i, run in enumerate(inline):
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
    
    def _convert_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """Convert docx to PDF using LibreOffice"""
        try:
            # Try LibreOffice conversion
            output_dir = os.path.dirname(pdf_path)
            result = subprocess.run([
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                docx_path
            ], capture_output=True, timeout=60)
            
            # LibreOffice creates PDF with same name as input
            expected_pdf = docx_path.replace('.docx', '.pdf')
            if os.path.exists(expected_pdf) and expected_pdf != pdf_path:
                shutil.move(expected_pdf, pdf_path)
            
            return os.path.exists(pdf_path)
        except subprocess.TimeoutExpired:
            logger.error("PDF conversion timed out")
            return False
        except FileNotFoundError:
            logger.warning("LibreOffice not found, trying alternative conversion")
            # Try docx2pdf as fallback (works on Windows/Mac)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                return os.path.exists(pdf_path)
            except Exception as e:
                logger.error(f"docx2pdf conversion failed: {str(e)}")
                return False
        except Exception as e:
            logger.error(f"PDF conversion error: {str(e)}")
            return False


def get_template_categories():
    """Get list of template categories"""
    return TEMPLATE_CATEGORIES


def get_placeholder_documentation():
    """Get documentation of available placeholders"""
    return {
        "personal": {
            "{{FULL_NAME}}": "Full name of the candidate",
            "{{EMAIL}}": "Email address",
            "{{PHONE}}": "Phone number",
            "{{ADDRESS}}": "Location/Address",
            "{{ID_NUMBER}}": "ID or Passport number",
            "{{LANGUAGES}}": "List of languages (comma-separated)",
            "{{LINKEDIN_URL}}": "LinkedIn profile URL",
            "{{LINKEDIN}}": "Same as {{LINKEDIN_URL}}",
            "{{PHOTO}}": "Profile photo placeholder"
        },
        "summary": {
            "{{SUMMARY}}": "Professional summary/profile",
            "{{PROFESSIONAL_SUMMARY}}": "Same as {{SUMMARY}}",
            "{{PROFILE}}": "Same as {{SUMMARY}}",
            "{{OBJECTIVE}}": "Career objective (same as summary)"
        },
        "skills": {
            "{{SKILLS}}": "Comma-separated list of skills",
            "{{SKILL_LIST}}": "Bullet-pointed list of skills"
        },
        "experience": {
            "{{EXPERIENCE_SECTION}}": "Full formatted experience section",
            "{{EXP_1_TITLE}}": "Job title for experience 1",
            "{{EXP_1_COMPANY}}": "Company name for experience 1",
            "{{EXP_1_DURATION}}": "Duration for experience 1",
            "{{EXP_1_DESCRIPTION}}": "Description for experience 1",
            "{{EXP_1_ACHIEVEMENTS}}": "Achievements for experience 1",
            "note": "Replace 1 with 2, 3, 4, or 5 for additional experiences"
        },
        "education": {
            "{{EDUCATION_SECTION}}": "Full formatted education section",
            "{{EDU_1_DEGREE}}": "Degree/qualification for education 1",
            "{{EDU_1_INSTITUTION}}": "Institution name for education 1",
            "{{EDU_1_YEAR}}": "Year/period for education 1",
            "note": "Replace 1 with 2 or 3 for additional education entries"
        },
        "references": {
            "{{REFERENCES_SECTION}}": "Full formatted references section",
            "{{REFERENCES}}": "Same as {{REFERENCES_SECTION}}",
            "{{REF_1_NAME}}": "Full name of reference 1",
            "{{REF_1_TITLE}}": "Job title of reference 1",
            "{{REF_1_COMPANY}}": "Company/organisation of reference 1",
            "{{REF_1_EMAIL}}": "Email address of reference 1",
            "{{REF_1_PHONE}}": "Phone number of reference 1",
            "note": "Replace 1 with 2 or 3 for additional references"
        },
        "certifications": {
            "{{CERTIFICATIONS_SECTION}}": "Full formatted certifications section",
            "{{CERTIFICATIONS}}": "Same as {{CERTIFICATIONS_SECTION}}",
            "{{CERT_1_NAME}}": "Name of certification 1",
            "{{CERT_1_ORGANIZATION}}": "Issuing organization for certification 1",
            "{{CERT_1_ISSUE_DATE}}": "Issue date for certification 1",
            "{{CERT_1_EXPIRY_DATE}}": "Expiry date for certification 1 (if applicable)",
            "{{CERT_1_CREDENTIAL_ID}}": "Credential ID for certification 1",
            "{{CERT_1_URL}}": "Verification URL for certification 1",
            "note": "Replace 1 with 2 or 3 for additional certifications"
        }
    }
